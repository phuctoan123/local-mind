from __future__ import annotations

from app.ingestion.parsers.base import BaseParser, ParsedPage


class DocxParser(BaseParser):
    def parse(self, file_path: str) -> list[ParsedPage]:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required to parse DOCX files") from exc

        doc = Document(file_path)
        blocks: list[str] = []
        current_heading: str | None = None
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if paragraph.style and paragraph.style.name.lower().startswith("heading"):
                current_heading = text
            prefix = f"{current_heading}\n" if current_heading and text != current_heading else ""
            blocks.append(prefix + text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))

        text = "\n\n".join(blocks)
        return _split_text_pages(text, page_size=3000)


def _split_text_pages(text: str, page_size: int) -> list[ParsedPage]:
    if not text.strip():
        return []
    pages = []
    for start in range(0, len(text), page_size):
        chunk = text[start : start + page_size].strip()
        if chunk:
            pages.append(
                ParsedPage(
                    page_number=len(pages) + 1,
                    text=chunk,
                    metadata={"logical_page": True},
                )
            )
    return pages
